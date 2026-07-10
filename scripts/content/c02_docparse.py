"""Einheitlicher Zugriff auf .docx und .doc: Bloecke in Dokument-Reihenfolge.

Ein Block ist ('para', text) oder ('table', [[zelle, ...], ...]).
Die Reihenfolge ist entscheidend: In den IQB-Auswertungen trennt der Absatz
'Teilaufgabe N' die Tabellen voneinander. Ohne Reihenfolge waere die Zuordnung
Loesung -> Teilaufgabe geraten.

.doc wird ueber antiword gelesen; dessen ASCII-Tabellen (|a|b|) werden
zurueckuebersetzt. Antiword bricht lange Zellen um, deshalb werden
Fortsetzungszeilen an die Vorgaengerzelle angehaengt.
"""
import re
import subprocess

PIPE_ROW = re.compile(r"^\s*\|.*\|\s*$")


def _docx_blocks(path):
    import docx
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = docx.Document(path)
    body = doc.element.body
    parts = {id(t._element): t for t in doc.tables}
    parts.update({id(p._element): p for p in doc.paragraphs})
    for child in body.iterchildren():
        obj = parts.get(id(child))
        if isinstance(obj, Paragraph):
            if obj.text.strip():
                yield "para", obj.text.strip()
        elif isinstance(obj, Table):
            yield "table", [[c.text.strip() for c in row.cells] for row in obj.rows]


def _split_pipe_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _doc_blocks(path):
    res = subprocess.run(["antiword", path], capture_output=True, text=True)
    if res.returncode != 0:
        return
    table = None
    for line in res.stdout.split("\n"):
        if PIPE_ROW.match(line):
            cells = _split_pipe_row(line)
            if table is None:
                table = [cells]
            elif _is_continuation(cells):
                _merge_continuation(table[-1], cells)
            else:
                table.append(cells)
            continue
        if not line.strip():
            continue  # Leerzeile trennt keine Tabelle - antiword setzt sie mitten hinein
        if table is not None:
            yield "table", table
            table = None
        yield "para", line.strip()
    if table is not None:
        yield "table", table


def _is_continuation(cells):
    """Fortsetzungszeile: entweder Wert-Umbruch (erste Zelle leer) oder
    Label-Umbruch (nur die erste Zelle traegt Text)."""
    if not cells:
        return False
    first_empty = not cells[0]
    rest_empty = not any(cells[1:])
    return first_empty or rest_empty


def _merge_continuation(prev, cells):
    if not cells[0]:  # Wert laeuft weiter
        for i in range(1, min(len(prev), len(cells))):
            if cells[i]:
                prev[i] = (prev[i] + "\n" + cells[i]).strip()
    else:  # Label laeuft weiter (antiword bricht 'Anforderungsber|eich' um)
        prev[0] = (prev[0] + cells[0]).strip()


def blocks(path):
    """[(kind, payload)] in Dokument-Reihenfolge. Leer bei Lesefehler."""
    if not path:
        return []
    try:
        if path.endswith(".docx"):
            return list(_docx_blocks(path))
        if path.endswith(".doc"):
            return list(_doc_blocks(path))
    except Exception:
        return []
    return []


def norm_ws(text):
    return re.sub(r"[^\S\n]+", " ", (text or "").replace("\xa0", " ")).strip()
