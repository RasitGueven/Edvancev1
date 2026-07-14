"""R01 — Stufe 1 (STRUKTUR) + Stufe 2 (SEHEN, Teil 1: rendern).

Erzeugt pro Item eine Quellen-Akte: was steht deterministisch in den DOCX, was
steht im EMF (Zeichenvorrat + Lage), und welche Bilder gehen an die Vision.

Diese Datei *interpretiert* nichts. Sie sammelt nur Belege. Die Auslegung
passiert in Stufe 2 (Vision) und wird in ground.py gegen genau diese Belege
geprueft.

    python3 extract.py                 # die fuenf Kalibrier-Items
"""
import hashlib
import json
import sys
from pathlib import Path

import docx

import layout
import table as tbl
from emf import parse, tokens
from render import emfs_in_docx, render_emf

ROOT = Path(__file__).resolve().parents[3]
DOCS = ROOT / "data/vera8_docs"
OUT = ROOT / "data/r01_extract"
RENDER = ROOT / "data/r01_render"

ITEMS = ["zeitangabe", "holzwuerfel", "kopfundkoerper", "bevoelkerungsdichte", "20prozent"]


def _sha(p):
    return hashlib.sha256(Path(p).read_bytes()).hexdigest()[:16]


def docx_struct(path):
    """Stufe 1: Absaetze und Tabellen ALS TABELLEN. Kein LLM, kein Plattwalzen."""
    d = docx.Document(path)
    return {
        "paragraphs": [p.text.strip() for p in d.paragraphs if p.text.strip()],
        "tables": [
            {"cells": [[c.text.strip() for c in r.cells] for r in t.rows]}
            for t in d.tables
        ],
    }


def aufgabe_struct(item):
    """Stufe 2: jedes EMF rendern; Zeichenvorrat und Lage deterministisch sichern."""
    path = DOCS / item / f"{item}_Aufgabe.docx"
    out = []
    for i, (name, data) in enumerate(emfs_in_docx(path)):
        p = parse(data)
        png = RENDER / item / f"{i:02d}_{name}.png"
        render_emf(data, png)
        out.append({
            "idx": i,
            "emf": name,
            "png": str(png.relative_to(ROOT)),
            # Der Zeichenvorrat ist die Grundwahrheit fuers Grounding: die
            # Vision darf die Reihenfolge bestimmen, aber keine Zeichen erfinden.
            "tokens": tokens(p),
            "deterministic_text": layout.text(p["texts"]),
            "table": tbl.extract(p),
            "n_fragments": len(p["texts"]),
        })
    return out


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    for item in ITEMS:
        d = DOCS / item
        rec = {
            "item": item,
            "sources": {
                k: {"file": f"{item}_{k}.docx", "sha256_16": _sha(d / f"{item}_{k}.docx")}
                for k in ("Aufgabe", "Auswertung", "Didaktische_Kommentierung")
                if (d / f"{item}_{k}.docx").exists()
            },
            "stufe1_auswertung": docx_struct(d / f"{item}_Auswertung.docx"),
            "stufe1_aufgabe_docx": docx_struct(d / f"{item}_Aufgabe.docx"),
            "stufe2_aufgabe_emf": aufgabe_struct(item),
        }
        (OUT / f"{item}.json").write_text(
            json.dumps(rec, ensure_ascii=False, indent=2), encoding="utf-8")
        n_tbl = sum(1 for e in rec["stufe2_aufgabe_emf"] if e["table"])
        print(f"{item:22s} EMFs={len(rec['stufe2_aufgabe_emf'])} "
              f"Tabellen(EMF)={n_tbl} "
              f"Auswertung-Tabellen={len(rec['stufe1_auswertung']['tables'])}")


if __name__ == "__main__":
    sys.exit(main())
