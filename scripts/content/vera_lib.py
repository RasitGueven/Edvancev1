"""Gemeinsame Helfer fuer den C02 Grounded Rebuild.

Kein Schreibzugriff hier drin - reine Lese-/Konvertierungs-Utilities, damit
Phase 0-4 dieselbe Slug-Aufloesung und Textnormalisierung benutzen.
"""
import json
import os
import re
import subprocess

DOCS_DIR = "data/vera8_docs"
ASSETS_DIR = "data/vera8_assets"
ENRICHED = "data/vera8_komplett_enriched.json"
MANIFEST = "data/vera8_assets_manifest.json"


def slug_of(item):
    """Ordnername in data/vera8_docs/ aus der Aufgaben-URL."""
    urls = item.get("iqb_urls") or item.get("urls") or {}
    url = urls.get("aufgabe") or ""
    if not url:
        return ""
    return re.sub(r"_Aufgabe\..*$", "", os.path.basename(url)).lower()


def folder_of(item):
    slug = slug_of(item)
    if not slug:
        return None
    path = os.path.join(DOCS_DIR, slug)
    return path if os.path.isdir(path) else None


def files_of(item):
    """{'aufgabe': path|None, 'auswertung': path|None, 'kommentierung': path|None}"""
    out = {"aufgabe": None, "auswertung": None, "kommentierung": None}
    folder = folder_of(item)
    if not folder:
        return out
    for fname in sorted(os.listdir(folder)):
        low = fname.lower()
        for key in out:
            if key in low and out[key] is None:
                out[key] = os.path.join(folder, fname)
    return out


def convert(path):
    """Rohtext aus .doc/.docx. None bei Fehler oder unbekanntem Format."""
    if path is None:
        return None
    if path.endswith(".doc"):
        res = subprocess.run(["antiword", path], capture_output=True, text=True)
        return res.stdout if res.returncode == 0 else None
    if path.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(path)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        except Exception:
            return None
    return None


def norm(text):
    """Fuer Grounding-Vergleiche: lower, Dezimalkomma->Punkt, Sonderzeichen raus."""
    t = (text or "").lower().replace(",", ".")
    t = t.replace(" ", " ")
    return re.sub(r"[^a-z0-9äöüß./%:-]+", "", t)


def load_items():
    return json.load(open(ENRICHED))


def save_items(items):
    with open(ENRICHED, "w") as fh:
        json.dump(items, fh, ensure_ascii=False, indent=1)
        fh.write("\n")
