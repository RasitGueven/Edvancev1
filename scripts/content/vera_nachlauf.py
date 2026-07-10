"""C01 Nachlauf - Phase 1: Beschaffung + Konvertierung der 81 Items ohne aufgabe_text_clean.

Quelle: lokaler Cache data/vera8_docs/<slug>/ ; fehlende Dateien werden hoeflich
von iqb.hu-berlin.de nachgeladen (sleep 1, User-Agent, 3x Retry mit Backoff).

Output:
  data/vera_nachlauf_raw/<slug>_{aufgabe,auswertung}.txt   - Rohtexte (Cache)
  data/vera_nachlauf_log.csv                               - Status je Item

Dieses Script STRUKTURIERT NICHT und schreibt NICHT in vera8_komplett.json.
Es stellt nur fest, was tatsaechlich an Text vorhanden ist.
"""
import csv
import json
import os
import re
import subprocess
import time
import urllib.error
import urllib.request

DOCS_DIR = "data/vera8_docs"
RAW_DIR = "data/vera_nachlauf_raw"
LOG_CSV = "data/vera_nachlauf_log.csv"
USER_AGENT = "EdvanceContentBot/1.0 (Bildungsprojekt Koeln; CC-BY-Nachnutzung; kontakt via repo)"
RETRIES = 3


def slug_of(item):
    url = item["urls"].get("aufgabe") or ""
    return re.sub(r"_Aufgabe\..*$", "", os.path.basename(url)).lower()


def fetch(url, dest):
    """Hoeflicher Download mit Backoff. True bei Erfolg."""
    for attempt in range(RETRIES):
        try:
            req = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = resp.read()
            os.makedirs(os.path.dirname(dest), exist_ok=True)
            with open(dest, "wb") as fh:
                fh.write(data)
            time.sleep(1)  # hoeflich crawlen
            return True
        except (urllib.error.URLError, urllib.error.HTTPError, OSError):
            time.sleep(2 ** attempt)
    return False


def convert(path):
    """Rohtext aus .doc/.docx. None bei Konvertierungsfehler."""
    if path.endswith(".doc"):
        res = subprocess.run(["antiword", path], capture_output=True, text=True)
        return res.stdout if res.returncode == 0 else None
    if path.endswith(".docx"):
        try:
            import docx
            doc = docx.Document(path)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        except Exception:
            return None
    return None


def text_chars(raw):
    """Zeichen echten Texts - ohne [pic]-Platzhalter, Tabellenrahmen, Whitespace."""
    if raw is None:
        return -1
    body = re.sub(r"\[pic\]", "", raw).replace("|", "")
    return len(re.sub(r"[\s_-]+", " ", body).strip())


def stem_chars(raw, titel):
    """Zeichen, die NACH Abzug von Titel + 'Teilaufgabe N:' + IQB-Item-IDs bleiben.

    Das ist der eigentliche Aufgabenstamm. Ist er ~0, steckt die Aufgabe im Bild.
    """
    if raw is None:
        return -1
    body = re.sub(r"\[pic\]", "", raw).replace("|", "")
    body = body.replace(titel, " ")
    body = re.sub(r"Teilaufgabe\s*\d+\s*:?", " ", body)
    body = re.sub(r"\bM\d{4,6}\b", " ", body)          # IQB-Item-IDs, z.B. M46017
    body = re.sub(r"^[A-Z0-9_]$", " ", body, flags=re.M)  # vertikal zerlegte IDs
    body = re.sub(r"[\s_\-A-Z0-9]+", " ", body)
    return len(body.strip())


def main():
    doc = json.load(open("data/vera8_komplett.json"))
    missing = [a for a in doc["aufgaben"] if not a.get("aufgabe_text_clean")]
    os.makedirs(RAW_DIR, exist_ok=True)

    rows = []
    for item in missing:
        titel, slug = item["iqb_titel"], slug_of(item)
        row = {"titel": titel, "slug": slug, "datei_ext": item["datei_ext"]}

        for kind in ("aufgabe", "auswertung"):
            url = item["urls"].get(kind)
            if not url:
                row[f"{kind}_status"] = "keine_url"
                row[f"{kind}_chars"] = -1
                continue
            if url.endswith(".ggb"):
                row[f"{kind}_status"] = "unsupported_ggb"
                row[f"{kind}_chars"] = -1
                continue

            local = os.path.join(DOCS_DIR, slug, os.path.basename(url))
            if not os.path.exists(local):
                if not fetch(url, local):
                    row[f"{kind}_status"] = "download_failed"
                    row[f"{kind}_chars"] = -1
                    continue

            raw = convert(local)
            if raw is None:
                row[f"{kind}_status"] = "convert_failed"
                row[f"{kind}_chars"] = -1
                continue

            with open(os.path.join(RAW_DIR, f"{slug}_{kind}.txt"), "w") as fh:
                fh.write(raw)
            row[f"{kind}_status"] = "ok"
            row[f"{kind}_chars"] = text_chars(raw)
            if kind == "aufgabe":
                row["aufgabe_stem_chars"] = stem_chars(raw, titel)

        # Verdikt: ist der Aufgabenstamm als TEXT vorhanden?
        stem = row.get("aufgabe_stem_chars", -1)
        if row.get("aufgabe_status") != "ok":
            row["verdikt"] = row.get("aufgabe_status", "kein_aufgabenfile")
        elif stem < 30:
            row["verdikt"] = "stem_nur_bild"      # Aufgabentext steckt im [pic]
        else:
            row["verdikt"] = "stem_text_vorhanden"
        rows.append(row)

    cols = ["titel", "slug", "datei_ext", "aufgabe_status", "aufgabe_chars",
            "aufgabe_stem_chars", "auswertung_status", "auswertung_chars", "verdikt"]
    with open(LOG_CSV, "w", newline="") as fh:
        w = csv.DictWriter(fh, fieldnames=cols, extrasaction="ignore")
        w.writeheader()
        for r in rows:
            w.writerow(r)

    from collections import Counter
    print(f"Items verarbeitet: {len(rows)}")
    print("\nVerdikt (kann aufgabe_text_clean befuellt werden?):")
    for k, v in Counter(r["verdikt"] for r in rows).most_common():
        print(f"  {k:24s} {v:3d}")
    print("\nAuswertungsdatei-Status:")
    for k, v in Counter(r.get("auswertung_status", "?") for r in rows).most_common():
        print(f"  {k:24s} {v:3d}")
    usable = sum(1 for r in rows if r.get("auswertung_chars", -1) > 50)
    print(f"\nAuswertungen mit verwertbarem Inhalt (>50 Zeichen): {usable}")
    print(f"Log: {LOG_CSV}   Rohtexte: {RAW_DIR}/")


if __name__ == "__main__":
    main()
