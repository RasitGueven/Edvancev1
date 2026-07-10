"""Diagnose-Lauf: Wieviel Text lässt sich aus den 81 Items überhaupt extrahieren?

Kein Schreibzugriff auf data/*.json - reine Messung.
"""
import json
import os
import re
import subprocess

ITEMS = json.load(open("data/vera8_komplett.json"))["aufgaben"]
MISSING = [a for a in ITEMS if not a.get("aufgabe_text_clean")]


def slug_of(item):
    url = item["urls"].get("aufgabe") or ""
    return re.sub(r"_Aufgabe\..*$", "", os.path.basename(url)).lower()


def convert(path):
    if path.endswith(".doc"):
        res = subprocess.run(["antiword", path], capture_output=True, text=True)
        return res.stdout if res.returncode == 0 else None
    if path.endswith(".docx"):
        try:
            import docx
        except ImportError:
            return None
        try:
            doc = docx.Document(path)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text for c in row.cells))
            return "\n".join(parts)
        except Exception:
            return None
    return None


def meat(text):
    """Zeichen echten Texts: [pic]-Platzhalter, Tabellenrahmen und Whitespace raus."""
    if text is None:
        return -1
    stripped = re.sub(r"\[pic\]", "", text)
    stripped = stripped.replace("|", "")
    stripped = re.sub(r"[\s_-]+", " ", stripped)
    return len(stripped.strip())


rows = []
for item in MISSING:
    slug = slug_of(item)
    folder = os.path.join("data/vera8_docs", slug)
    if not os.path.isdir(folder):
        rows.append({"titel": item["iqb_titel"], "slug": slug, "status": "no_folder",
                     "aufg_chars": -1, "ausw_chars": -1})
        continue
    files = os.listdir(folder)
    aufg = next((f for f in files if "Aufgabe" in f), None)
    ausw = next((f for f in files if "Auswertung" in f), None)
    rows.append({
        "titel": item["iqb_titel"],
        "slug": slug,
        "status": "ok",
        "aufg_file": aufg,
        "aufg_chars": meat(convert(os.path.join(folder, aufg))) if aufg else -1,
        "ausw_chars": meat(convert(os.path.join(folder, ausw))) if ausw else -1,
    })

json.dump(rows, open("/tmp/census.json", "w"), ensure_ascii=False, indent=1)

print(f"total: {len(rows)}")
for thr in (0, 50, 100, 200, 400):
    print(f"  aufgabe-Text > {thr:4d} Zeichen: {sum(1 for r in rows if r['aufg_chars'] > thr):3d}")
print(f"  auswertung   >   50 Zeichen: {sum(1 for r in rows if r['ausw_chars'] > 50):3d}")

print("\n--- 15 niedrigste (aufgabe_chars) ---")
for r in sorted(rows, key=lambda x: x["aufg_chars"])[:15]:
    print(f"{r['aufg_chars']:6d} {r['ausw_chars']:6d}  {r['titel'][:42]}")
print("\n--- 10 höchste ---")
for r in sorted(rows, key=lambda x: -x["aufg_chars"])[:10]:
    print(f"{r['aufg_chars']:6d} {r['ausw_chars']:6d}  {r['titel'][:42]}")
